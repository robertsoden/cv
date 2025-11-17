#!/usr/bin/env python3
"""
Generate CV documents from templates using parsed CV data and Google Scholar data.
"""

import json
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CVGenerator:
    """Generator for creating CV documents from templates and data."""

    def __init__(self, cv_data_path: str = 'data/cv_data.json',
                 publications_path: str = 'data/publications.json',
                 include_citations: bool = False):
        """
        Initialize the CV generator.

        Args:
            cv_data_path: Path to parsed CV data JSON
            publications_path: Path to Google Scholar publications JSON
            include_citations: Whether to include citation counts (default: False)
        """
        self.cv_data = self._load_json(cv_data_path)
        self.publications_data = self._load_json(publications_path)
        self.include_citations = include_citations

    def _load_json(self, path: str) -> Dict:
        """Load JSON data from file."""
        path = Path(path)
        if not path.exists():
            print(f"Warning: {path} not found. Using empty data.")
            return {}

        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def get_merged_data(self) -> Dict[str, Any]:
        """
        Merge CV data with publications data.

        Returns:
            Dictionary with all merged CV data
        """
        merged = {
            'personal_info': self.cv_data.get('personal_info', {}),
            'education': self.cv_data.get('education', []),
            'experience': self.cv_data.get('experience', []),
            'publications': self.publications_data.get('publications', []),
            'author_info': self.publications_data.get('author_info', {}),
            'awards': self.cv_data.get('awards', []),
            'skills': self.cv_data.get('skills', []),
            'other_sections': self.cv_data.get('other_sections', {}),
            'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }

        return merged

    def generate_full_academic_cv(self, output_path: str = 'output/academic_cv_full.docx'):
        """
        Generate a full academic CV with all details.

        Args:
            output_path: Path to save the generated CV
        """
        print("Generating full academic CV...")

        doc = Document()
        merged_data = self.get_merged_data()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Header - Name
        personal_info = merged_data.get('personal_info', {})
        name = personal_info.get('name', 'Your Name')
        heading = doc.add_heading(name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact Information
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(f"Email: {personal_info['email']}")
        if personal_info.get('phone'):
            contact_info.append(f"Phone: {personal_info['phone']}")

        author_info = merged_data.get('author_info', {})
        if author_info.get('affiliation'):
            contact_info.append(author_info['affiliation'])

        contact_para.add_run(' | '.join(contact_info))

        # Google Scholar Metrics
        if author_info:
            doc.add_heading('Research Metrics', level=2)
            metrics_para = doc.add_paragraph()
            metrics_text = (
                f"Citations: {author_info.get('citedby', 0)} | "
                f"h-index: {author_info.get('hindex', 0)} | "
                f"i10-index: {author_info.get('i10index', 0)}"
            )
            metrics_para.add_run(metrics_text)

            if author_info.get('url'):
                metrics_para.add_run(f"\nGoogle Scholar Profile: {author_info['url']}")

        # Research Interests
        if author_info.get('interests'):
            doc.add_heading('Research Interests', level=2)
            interests_para = doc.add_paragraph()
            interests_para.add_run(', '.join(author_info['interests']))

        # Publications
        publications = merged_data.get('publications', [])
        if publications:
            doc.add_heading('Publications', level=2)

            # Sort by year (descending) and citations
            sorted_pubs = sorted(publications,
                               key=lambda x: (int(x.get('year', 0) or 0), x.get('citations', 0)),
                               reverse=True)

            for idx, pub in enumerate(sorted_pubs, 1):
                pub_para = doc.add_paragraph(style='List Number')

                # Format: Authors (Year). Title. Venue. [Citations]
                authors = pub.get('authors', '')
                year = pub.get('year', '')
                title = pub.get('title', '')
                venue = pub.get('venue', '')
                citations = pub.get('citations', 0)

                citation_text = f"{authors} ({year}). {title}."
                if venue:
                    citation_text += f" {venue}."

                pub_para.add_run(citation_text)

                # Add citation count (only if enabled)
                if self.include_citations and citations > 0:
                    cite_run = pub_para.add_run(f" [Cited by {citations}]")
                    cite_run.font.italic = True
                    cite_run.font.color.rgb = RGBColor(100, 100, 100)

        # Education
        if merged_data.get('other_sections', {}).get('education'):
            doc.add_heading('Education', level=2)
            for item in merged_data['other_sections']['education']:
                doc.add_paragraph(item, style='List Bullet')

        # Experience
        if merged_data.get('other_sections', {}).get('experience'):
            doc.add_heading('Experience', level=2)
            for item in merged_data['other_sections']['experience']:
                doc.add_paragraph(item, style='List Bullet')

        # Awards and Honors
        if merged_data.get('other_sections', {}).get('awards'):
            doc.add_heading('Awards and Honors', level=2)
            for item in merged_data['other_sections']['awards']:
                doc.add_paragraph(item, style='List Bullet')

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        print(f"Full academic CV saved to: {output_path}")
        return output_path

    def generate_short_cv(self, output_path: str = 'output/academic_cv_short.docx',
                         max_publications: int = 10):
        """
        Generate a short/abbreviated CV.

        Args:
            output_path: Path to save the generated CV
            max_publications: Maximum number of publications to include
        """
        print("Generating short academic CV...")

        doc = Document()
        merged_data = self.get_merged_data()

        # Header - Name
        personal_info = merged_data.get('personal_info', {})
        name = personal_info.get('name', 'Your Name')
        heading = doc.add_heading(name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact Information
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_info = []
        if personal_info.get('email'):
            contact_info.append(personal_info['email'])

        author_info = merged_data.get('author_info', {})
        if author_info.get('affiliation'):
            contact_info.append(author_info['affiliation'])

        contact_para.add_run(' | '.join(contact_info))

        # Research Summary
        if author_info.get('interests'):
            doc.add_heading('Research Interests', level=2)
            interests_para = doc.add_paragraph()
            interests_para.add_run(', '.join(author_info['interests']))

        # Selected Publications
        publications = merged_data.get('publications', [])
        if publications:
            doc.add_heading(f'Selected Publications (Top {max_publications})', level=2)

            # Sort by citations
            sorted_pubs = sorted(publications,
                               key=lambda x: x.get('citations', 0),
                               reverse=True)[:max_publications]

            for idx, pub in enumerate(sorted_pubs, 1):
                pub_para = doc.add_paragraph(style='List Number')

                authors = pub.get('authors', '')
                year = pub.get('year', '')
                title = pub.get('title', '')
                venue = pub.get('venue', '')
                citations = pub.get('citations', 0)

                citation_text = f"{authors} ({year}). {title}."
                if venue:
                    citation_text += f" {venue}."

                pub_para.add_run(citation_text)

                if self.include_citations and citations > 0:
                    cite_run = pub_para.add_run(f" [{citations} citations]")
                    cite_run.font.italic = True

        # Education (abbreviated)
        if merged_data.get('other_sections', {}).get('education'):
            doc.add_heading('Education', level=2)
            # Take only first 3 items
            for item in merged_data['other_sections']['education'][:3]:
                doc.add_paragraph(item, style='List Bullet')

        # Save document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        print(f"Short CV saved to: {output_path}")
        return output_path

    def export_publications_list(self, output_path: str = 'output/publications_list.txt'):
        """
        Export a simple text list of all publications.

        Args:
            output_path: Path to save the publications list
        """
        print("Generating publications list...")

        publications = self.publications_data.get('publications', [])

        if not publications:
            print("No publications found.")
            return

        # Sort by year
        sorted_pubs = sorted(publications,
                           key=lambda x: int(x.get('year', 0) or 0),
                           reverse=True)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("PUBLICATIONS LIST\n")
            f.write("=" * 80 + "\n\n")

            author_info = self.publications_data.get('author_info', {})
            if author_info:
                f.write(f"Author: {author_info.get('name', '')}\n")
                f.write(f"Total citations: {author_info.get('citedby', 0)}\n")
                f.write(f"h-index: {author_info.get('hindex', 0)}\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d')}\n\n")
                f.write("=" * 80 + "\n\n")

            for idx, pub in enumerate(sorted_pubs, 1):
                f.write(f"{idx}. {pub.get('authors', '')} ({pub.get('year', '')}). "
                       f"{pub.get('title', '')}. {pub.get('venue', '')}.\n")
                if pub.get('citations', 0) > 0:
                    f.write(f"   Citations: {pub['citations']}\n")
                f.write("\n")

        print(f"Publications list saved to: {output_path}")
        return output_path


def main():
    """Main function to generate CVs."""
    import sys

    # Parse command line arguments
    include_citations = '--include-citations' in sys.argv or '-c' in sys.argv

    print("=" * 60)
    print("Academic CV Generator")
    print("=" * 60)
    if include_citations:
        print("Citation counts: ENABLED")
    else:
        print("Citation counts: DISABLED (use --include-citations to enable)")
    print("=" * 60)

    generator = CVGenerator(include_citations=include_citations)

    # Check if data files exist
    if not generator.cv_data and not generator.publications_data:
        print("\nError: No data found!")
        print("Please run the following first:")
        print("  1. python scripts/parse_cv.py source_cv/your_cv.docx")
        print("  2. python scripts/fetch_scholar.py YOUR_SCHOLAR_ID")
        sys.exit(1)

    # Generate all CV formats
    print("\nGenerating CV documents...\n")

    generator.generate_full_academic_cv()
    generator.generate_short_cv()
    generator.export_publications_list()

    print("\n" + "=" * 60)
    print("All CV documents generated successfully!")
    print("Check the 'output/' directory for generated files.")
    print("=" * 60)


if __name__ == '__main__':
    main()
