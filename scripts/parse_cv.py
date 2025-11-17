#!/usr/bin/env python3
"""
Parse CV from Word document and extract structured data.
"""

import json
import re
from pathlib import Path
from typing import Dict, List, Any
from docx import Document


class CVParser:
    """Parser for extracting structured data from CV Word documents."""

    def __init__(self, docx_path: str):
        """
        Initialize the CV parser.

        Args:
            docx_path: Path to the Word document CV
        """
        self.docx_path = Path(docx_path)
        self.document = Document(docx_path)
        self.cv_data = {
            'personal_info': {},
            'education': [],
            'experience': [],
            'publications': [],
            'awards': [],
            'skills': [],
            'other_sections': {}
        }

    def extract_text_by_paragraphs(self) -> List[str]:
        """Extract all paragraphs from the document."""
        return [para.text.strip() for para in self.document.paragraphs if para.text.strip()]

    def extract_tables(self) -> List[List[List[str]]]:
        """Extract all tables from the document."""
        tables_data = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        return tables_data

    def identify_section(self, text: str) -> str:
        """
        Identify which section a text belongs to based on common CV section headers.

        Args:
            text: Text to analyze

        Returns:
            Section identifier or 'unknown'
        """
        text_lower = text.lower()

        # Common section headers
        if any(keyword in text_lower for keyword in ['education', 'academic background', 'degrees']):
            return 'education'
        elif any(keyword in text_lower for keyword in ['experience', 'employment', 'positions', 'appointments']):
            return 'experience'
        elif any(keyword in text_lower for keyword in ['publication', 'research output', 'papers']):
            return 'publications'
        elif any(keyword in text_lower for keyword in ['award', 'honor', 'recognition', 'grant', 'funding']):
            return 'awards'
        elif any(keyword in text_lower for keyword in ['skill', 'competenc', 'expertise']):
            return 'skills'
        elif any(keyword in text_lower for keyword in ['contact', 'email', 'phone', 'address']):
            return 'personal_info'

        return 'unknown'

    def extract_personal_info(self, paragraphs: List[str]) -> Dict[str, str]:
        """
        Extract personal information from the document.

        Args:
            paragraphs: List of paragraph texts

        Returns:
            Dictionary with personal information
        """
        personal_info = {}

        # Typically name is in the first paragraph (often in larger font)
        if paragraphs:
            personal_info['name'] = paragraphs[0]

        # Look for email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        for para in paragraphs[:10]:  # Check first 10 paragraphs
            email_match = re.search(email_pattern, para)
            if email_match:
                personal_info['email'] = email_match.group()
                break

        # Look for phone
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        for para in paragraphs[:10]:
            phone_match = re.search(phone_pattern, para)
            if phone_match:
                personal_info['phone'] = phone_match.group()
                break

        return personal_info

    def parse(self) -> Dict[str, Any]:
        """
        Parse the CV document and extract structured data.

        Returns:
            Dictionary with structured CV data
        """
        paragraphs = self.extract_text_by_paragraphs()
        tables = self.extract_tables()

        # Extract personal information
        self.cv_data['personal_info'] = self.extract_personal_info(paragraphs)

        # Group paragraphs by section
        current_section = 'unknown'
        section_content = []

        for para in paragraphs:
            section_type = self.identify_section(para)

            # Check if this is a section header
            if section_type != 'unknown' and len(para) < 100:
                # Save previous section
                if current_section != 'unknown' and section_content:
                    self.cv_data['other_sections'][current_section] = section_content

                current_section = section_type
                section_content = []
            else:
                section_content.append(para)

        # Save last section
        if current_section != 'unknown' and section_content:
            self.cv_data['other_sections'][current_section] = section_content

        # Store all paragraphs for reference
        self.cv_data['all_paragraphs'] = paragraphs

        # Store all tables
        self.cv_data['tables'] = tables

        return self.cv_data

    def save_to_json(self, output_path: str = None):
        """
        Save parsed CV data to JSON file.

        Args:
            output_path: Path to save JSON file (default: data/cv_data.json)
        """
        if output_path is None:
            output_path = 'data/cv_data.json'

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.cv_data, f, indent=2, ensure_ascii=False)

        print(f"CV data saved to {output_path}")


def main():
    """Main function to parse CV."""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parse_cv.py <path_to_cv.docx>")
        print("Example: python parse_cv.py source_cv/my_cv.docx")
        sys.exit(1)

    cv_path = sys.argv[1]

    if not Path(cv_path).exists():
        print(f"Error: File not found: {cv_path}")
        sys.exit(1)

    print(f"Parsing CV from: {cv_path}")

    parser = CVParser(cv_path)
    cv_data = parser.parse()

    # Save to JSON
    parser.save_to_json()

    # Print summary
    print("\n=== CV Parsing Summary ===")
    print(f"Name: {cv_data['personal_info'].get('name', 'Not found')}")
    print(f"Email: {cv_data['personal_info'].get('email', 'Not found')}")
    print(f"Total paragraphs: {len(cv_data.get('all_paragraphs', []))}")
    print(f"Total tables: {len(cv_data.get('tables', []))}")
    print(f"Sections found: {', '.join(cv_data['other_sections'].keys())}")


if __name__ == '__main__':
    main()
